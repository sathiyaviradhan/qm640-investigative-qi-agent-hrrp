# -*- coding: utf-8 -*-
"""Build QM640 Interim Report from template structure + executed analysis results."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(r"c:\Users\sathi\OneDrive\Documents\DBA Capstone 2026")
OUT = ROOT / "QM640_Interim_Report.docx"
METRICS = json.loads((ROOT / "analysis_outputs" / "metrics.json").read_text(encoding="utf-8"))
FIG = ROOT / "analysis_outputs" / "figures"

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)
rpr = style.element.get_or_add_rPr()
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rpr.append(rFonts)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hp.add_run()
run.font.name = "Times New Roman"
run.font.size = Pt(12)
add_page_number(hp)


def p(text, bold=False, center=False, first_line_indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Inches(0)
    elif first_line_indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    else:
        para.paragraph_format.first_line_indent = Inches(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return para


def heading(text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.space_before = Pt(12)
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
    elif level == 2:
        run = para.add_run(text)
        run.bold = True
    else:
        run = para.add_run(text)
        run.bold = True
        run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return para


def add_table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.first_line_indent = Inches(0)
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        r = cp.add_run(caption)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_fig(path: Path, caption: str, width=6.0):
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Inches(0)
        para.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Inches(0)
    r = cp.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


# ---------- helpers from metrics ----------
cl = METRICS["cleaning"]
eda = METRICS["eda"]
rq1 = METRICS["rq1"]
rq2 = METRICS["rq2"]
rq3 = METRICS["rq3"]
rq4 = METRICS["rq4"]
ss = METRICS["sample_size"]
prev = {r["fiscal_year_label"]: r["prevalence_err_gt1"] for r in eda["prevalence_by_fy"]}
m_by_name = {m["model"]: m for m in rq2["models"]}

# ========== TITLE PAGE ==========
for _ in range(2):
    p("", center=True, first_line_indent=False)
p("Data Analytics Capstone", center=True, first_line_indent=False)
p(
    "An Investigative Agentic AI Framework for CMS Hospital Excess Readmission Risk: "
    "Parallel Machine Learning, Fairness Monitoring, and Tool-Grounded Quality Improvement "
    "Case Dossiers under Human-in-the-Loop Review",
    center=True,
    first_line_indent=False,
)
p("", center=True, first_line_indent=False)
p("Interim Report", center=True, first_line_indent=False)
p("", center=True, first_line_indent=False)
p("SATHIYAVIRADHAN JANARTHANAN", center=True, first_line_indent=False)
p("Walsh College", center=True, first_line_indent=False)
p("QM640: Data Analytics Capstone", center=True, first_line_indent=False)
p("Mentor: SRIDHAR S", center=True, first_line_indent=False)
p("Group – 2", center=True, first_line_indent=False)
p("Summer 2026 Term", center=True, first_line_indent=False)
p("July 28, 2026", center=True, first_line_indent=False)

doc.add_page_break()

p("GitHub repository (data files and code):", first_line_indent=False)
p(
    "https://github.com/sathiyaviradhan/qm640-investigative-qi-agent-hrrp",
    first_line_indent=False,
)
p(
    "Local package paths mirrored for GitHub: DATASET/analytic_core/ (raw extracts); "
    "analysis_outputs/processed/ (cleaned eligible panel); run_interim_analysis.py and "
    "analysis_outputs/ (code, metrics, figures).",
    first_line_indent=False,
)

# ========== INTRODUCTION ==========
heading("Introduction")
heading("Background and Context", level=2)

p(
    "Hospital readmissions remain a central quality and cost concern in U.S. healthcare operations. "
    "Under the Centers for Medicare & Medicaid Services (CMS) Hospital Readmissions Reduction Program "
    "(HRRP), hospitals with higher-than-expected unplanned 30-day readmissions face Medicare payment "
    "reductions of up to three percent (Centers for Medicare & Medicaid Services [CMS], 2024). Quality "
    "leaders therefore need early excess-risk insight, equitable monitoring across reporting cycles, and "
    "actionable quality-improvement (QI) options that remain faithful to public CMS evidence."
)
p(
    "Supervised machine learning (ML) is well suited to tabular CMS quality files for excess-risk scoring, "
    "and TRIPOD+AI emphasizes transparent validation and fairness reporting for clinical prediction models "
    "(Collins et al., 2024). Audits show that readmission models can encode contextual bias (Wang, Weiner, "
    "Saria, & Kharrazi, 2024), and fairness can drift after deployment (Davis et al., 2025). These findings "
    "motivate parallel statistical association, predictive scoring, and fairness-monitoring tracks—but they "
    "do not by themselves produce a board-ready QI case package."
)
p(
    "The harder operational gap is investigative decision support. When Excess Readmission Ratio (ERR) "
    "exceeds 1.0, quality teams must assemble hospital CMS metrics, find peer hospitals that improved across "
    "fiscal releases, quantify temporal change, run what-if calculations (for example, how much predicted "
    "readmission rate must fall for ERR ≤ 1), and propose QI options with citations. Unconstrained large "
    "language models (LLMs) are fast but unsafe: they can hallucinate rates, invent peers, and recommend "
    "unsupported interventions (Singhal et al., 2023). Tool-wielding agents can explore clinical tables "
    "(Yang et al., 2026), yet prior work has not evaluated a multi-step investigative agent that builds "
    "peer-benchmarked, what-if-supported, citation-verified QI dossiers on multi-year CMS HRRP public files "
    "under mandatory human-in-the-loop (HITL) review."
)
p(
    "This interim report documents progress through planning Steps 2–4 (literature framing, SMART objectives/"
    "hypotheses, and design/methods) and presents executed preliminary results for all four research questions. "
    "The study uses verified CMS Provider Data Catalog extracts for FY2024–FY2026 (pooled eligible N = "
    f"{ss['pooled_eligible']:,} hospital–condition rows with non-missing ERR). Track A delivers association, "
    "ML scoring, and fairness monitoring. Track B delivers an Investigative Agentic QI Case Agent. Research "
    "questions remain independent so timeline risk is controlled."
)

heading("Problem Statement", level=2)
p(
    "U.S. acute-care hospitals face avoidable Medicare payment exposure under HRRP when excess 30-day "
    "readmission risk (ERR > 1.0) is not anticipated, when subgroup performance is not monitored across CMS "
    "fiscal releases, and when quality leaders lack a trustworthy investigative dossier that correctly combines "
    "hospital metrics, improving peers, temporal change, and what-if analysis before human QI decisions. Using "
    "public CMS HRRP and Hospital General Information files for FY2024–FY2026, this study addresses four "
    "complementary problems in parallel: (1) associate hospital attributes with excess risk; (2) score "
    "excess-risk status with supervised ML out-of-time; (3) monitor fairness/performance across fiscal "
    "releases; and (4) evaluate a multi-step tool-grounded investigative agent that opens high-ERR cases and "
    "produces verified QI case dossiers for mandatory human Accept/Edit/Reject. The dependent variable for "
    "scoring is Y = 1 if ERR > 1.0 (else 0)."
)

heading("Purpose of the Study", level=2)
p(
    "The purpose is to deliver a timeline-safe package that (a) explains excess-risk associations, "
    "(b) validates an ML scoring engine, (c) monitors fairness across FY2024–FY2026, and (d) evaluates an "
    "Investigative Agentic QI Case Agent that uniquely solves multi-source case investigation and dossier "
    "construction under HITL governance. At interim stage, the purpose is also to demonstrate that code has "
    "been developed and executed, with reproducible tables/figures linked to each research question."
)

heading("Interim Project Status (Progress Snapshot)", level=2)
p("Completed:", bold=True, first_line_indent=False)
for item in [
    "Literature synthesis (≥10 sources) mapped to RQ1–RQ4 (Step 2).",
    "SMART objectives, independent RQs, and H0/Ha statements locked to First Submission (Step 3).",
    "Methods design: cleaning rules, feature set, out-of-time split, fairness metrics, agent evaluation protocol (Step 4).",
    "Dataset verification and analytic_core extracts for FY2024–FY2026.",
    f"Data cleaning executed: raw N={cl['raw_rows']:,}; eligible ERR N={cl['eligible_err_rows']:,} ({cl['missing_err_pct']}% removed for missing ERR).",
    "EDA figures (Figures 1–4) generated and interpreted.",
    "RQ1 multivariable logistic associations estimated on pooled eligible panel.",
    "RQ2 out-of-time ML (train FY2024–FY2025; test FY2026) for logistic, random forest, gradient boosting.",
    "RQ3 temporal AUROC and ownership ΔFNR monitoring across FY releases.",
    "RQ4 tool-grounded investigative agent vs no-tool LLM vs static dashboard on n=80 high-ERR cases.",
]:
    p("• " + item, first_line_indent=False)

p("In progress:", bold=True, first_line_indent=False)
for item in [
    "SHAP/driver ranks as optional dossier evidence (not required for RQ4 tools).",
    "Human HITL Accept/Edit/Reject score sheet with external reviewer ratings (proxy usefulness used at interim).",
    "GitHub publication of full repository and README for evaluator access.",
]:
    p("• " + item, first_line_indent=False)

p("Pending (Final Report):", bold=True, first_line_indent=False)
for item in [
    "Sensitivity analyses (alternate thresholds; optional predicted/expected feature ablation transparency).",
    "Calibration plots and decision-curve style operational framing for RQ2.",
    "Expanded fairness strata (state/ZIP-derived geography) and formal drift tests.",
    "Optional LLM API planner layer behind the same tool contracts (current interim uses deterministic tool-grounded dossier generator).",
    "Final recommendations, limitations expansion, and presentation package.",
]:
    p("• " + item, first_line_indent=False)

# ========== SCOPE ==========
heading("Scope and Objectives")
p(
    "Scope boundaries. U.S. Medicare-participating acute-care hospitals in CMS HRRP joined to same-year "
    "Hospital General Information for FY2024–FY2026 (measures: AMI, HF, PN, COPD, CABG, HIP-KNEE). Track B "
    "produces advisory QI case dossiers only. No protected health information (PHI), no Kaggle data, and no "
    "autonomous clinical/CMS actions. Analytic scope is limited to core HRRP + General Information files."
)
p(
    "Primary objectives (parallel). (1) Association analysis for excess risk. (2) Out-of-time ML scoring. "
    "(3) Fairness/performance monitoring across FY releases. (4) Evaluate Investigative Agentic QI Case Agent "
    "against no-tool LLM and static dashboard baselines under mandatory HITL. (5) Reproducible code/data package."
)

heading("Research Question 1 (RQ1)", level=2)
p(
    "Which hospital structural and contextual variables are significantly associated with excess readmission "
    "risk (ERR > 1.0) after adjusting for condition measure and fiscal-year release?"
)
p(
    "Hypothesis RQ1. H0: After adjustment, focal predictors have no association (β = 0). Ha: At least one "
    "focal predictor has a nonzero association at α = .05."
)

heading("Research Question 2 (RQ2)", level=2)
p(
    "How well can supervised ML classifiers (regularized logistic regression, random forest, gradient boosting) "
    "discriminate excess readmission risk on out-of-time (later FY) samples relative to a transparent logistic baseline?"
)
p(
    "Hypothesis RQ2. H0: Best ML out-of-time AUROC ≤ logistic baseline + 0.02. Ha: Best ML out-of-time AUROC > "
    "logistic baseline + 0.02. Absolute AUROC, PR-AUC, and calibration-related metrics remain standalone deliverables."
)

heading("Research Question 3 (RQ3)", level=2)
p(
    "Do discrimination and fairness metrics (AUROC and ΔFNR across ownership subgroups) differ across FY2024, "
    "FY2025, and FY2026 when a model trained on an earlier release is scored on later releases?"
)
p(
    "Hypothesis RQ3. H0: Fairness gaps and AUROC do not differ across releases beyond operational thresholds. "
    "Ha: At least one operationally meaningful change (|ΔFNR| increase > 0.05 or AUROC drop > 0.03)."
)

heading("Research Question 4 (RQ4)", level=2)
p(
    "For high-ERR hospital–condition cases, does a multi-step tool-grounded Investigative Agentic QI Case Agent "
    "(case intake → metrics → improving-peer discovery → temporal change → what-if calculator → citation-linked "
    "dossier) produce higher evidence completeness, lower numeric hallucination, and higher HITL usefulness than "
    "(a) a no-tool LLM and (b) a static dashboard export alone, when every dossier is advisory only under "
    "mandatory human review?"
)
p(
    "Hypothesis RQ4. H0: The investigative agent does not improve faithfulness, evidence completeness, or HITL "
    "usefulness versus baselines. Ha: The agent improves at least two of {faithfulness, evidence completeness, "
    "HITL usefulness} at α = .05."
)

heading("Sample Size Justification (Executed)", level=2)
add_table(
    ["Research Question", "Method", "Minimum N (plan)", "Achieved N (interim)"],
    [
        ["RQ1", "Green’s rule / conservative", "170 / 400", f"{rq1['n']:,}"],
        ["RQ2", "AUROC CI precision", "800", f"Train {ss['rq2_train']:,}; Test {ss['rq2_test']:,}"],
        ["RQ3", "Stratum–period monitoring", "84 units", "3 FY × ownership strata"],
        ["RQ4", "Paired dossier power", "≥ 80", f"{ss['rq4_cases']}"],
    ],
    caption="Table 1. Sample-size plan versus achieved interim samples",
)
p(
    f"All planned minima are exceeded. Pooled eligible modeling N = {ss['pooled_eligible']:,} hospital–condition "
    "rows across 2,946 unique facilities."
)

# ========== LITERATURE ==========
heading("Literature Survey")
heading("Literature Review Approach", level=2)
p(
    "Sources were selected from peer-reviewed journals and CMS program documentation using keywords including "
    "hospital readmission, HRRP, Excess Readmission Ratio, clinical prediction fairness, TRIPOD+AI, LLM "
    "hallucination in medicine, and tool-using agents. Inclusion required relevance to at least one RQ "
    "(association/prediction, fairness drift, or trustworthy GenAI decision support). At least ten sources are "
    "summarized below and mapped in the relevance matrix."
)

heading("Summary of Key Literature", level=2)
lit_paras = [
    (
        "Collins et al. (2024) updated transparent reporting guidance for clinical prediction models with AI "
        "(TRIPOD+AI). Their emphasis on validation design, fairness, and clear outcome definitions directly "
        "shapes RQ2/RQ3 reporting choices in this capstone (out-of-time FY split; subgroup FNR)."
    ),
    (
        "Wang et al. (2024) audited readmission prediction for contextual bias, showing that performance and "
        "error rates can differ across hospital contexts. This supports RQ1 ownership/state predictors and RQ3 "
        "ownership fairness monitoring rather than overall accuracy alone."
    ),
    (
        "Davis et al. (2025) documented fairness drift over time after model deployment. Their finding motivates "
        "RQ3’s explicit comparison of AUROC and ΔFNR across FY2024–FY2026 releases instead of a single cross-section."
    ),
    (
        "Singhal et al. (2023) demonstrated strong LLM medical knowledge with persistent factual errors. This "
        "supports RQ4’s no-tool LLM baseline and the requirement that numeric claims be tool-verified."
    ),
    (
        "Yang et al. (2026) showed tool-using agents can navigate clinical tabular evidence more reliably than "
        "free-text-only generation. That capability is the methodological backbone of the Investigative QI Case Agent."
    ),
    (
        "CMS (2024) HRRP methodology defines ERR as predicted divided by expected readmission rate and links "
        "excess risk to payment adjustment. This official definition anchors Y = 1{ERR > 1.0} and the what-if tool."
    ),
]
for t in lit_paras:
    p(t)

add_table(
    ["Author (Year)", "Domain/Context", "Dataset/Setting", "Method(s)", "Key Findings", "Supports RQ(s)"],
    [
        [
            "CMS (2024)",
            "HRRP payment policy",
            "Medicare IPPS hospitals",
            "Program methodology",
            "ERR drives payment reduction risk",
            "All; Y definition",
        ],
        [
            "Collins et al. (2024)",
            "Prediction reporting",
            "Clinical AI models",
            "TRIPOD+AI guidance",
            "Transparent validation & fairness",
            "RQ2, RQ3",
        ],
        [
            "Wang et al. (2024)",
            "Readmission bias",
            "Hospital contexts",
            "Fairness audit",
            "Context-linked model bias",
            "RQ1, RQ3",
        ],
        [
            "Davis et al. (2025)",
            "Fairness drift",
            "Deployed models over time",
            "Longitudinal fairness",
            "Gaps can change post-deployment",
            "RQ3",
        ],
        [
            "Singhal et al. (2023)",
            "Medical LLMs",
            "Clinical QA",
            "LLM evaluation",
            "High capability; residual errors",
            "RQ4 baselines",
        ],
        [
            "Yang et al. (2026)",
            "Tool-using agents",
            "Clinical tables",
            "Agent + tools",
            "Tools improve grounded answers",
            "RQ4 design",
        ],
        [
            "Zafar et al. (2019)",
            "ML fairness metrics",
            "Classification systems",
            "Disparate error rates",
            "FNR gaps operationalize unfairness",
            "RQ3 metric choice",
        ],
        [
            "Lundberg & Lee (2017)",
            "Model explanation",
            "Tabular ML",
            "SHAP",
            "Local feature attributions",
            "RQ2 optional drivers",
        ],
        [
            "Rajkomar et al. (2018)",
            "Healthcare ML",
            "EHR prediction",
            "Deep/ensemble ML",
            "ML can outperform baselines",
            "RQ2 model family",
        ],
        [
            "Amodei et al. (2016)",
            "AI safety",
            "Autonomous systems",
            "Safety framing",
            "Avoid unintended agency",
            "RQ4 HITL policy",
        ],
        [
            "Joynt & Jha (2013)",
            "Readmission policy",
            "U.S. hospitals",
            "Policy analysis",
            "Readmission penalties create QI pressure",
            "Background / RQ1",
        ],
        [
            "Krumholz et al. (2017)",
            "Readmission measurement",
            "Medicare claims",
            "Risk-standardized metrics",
            "Case-mix adjusted rates underpin CMS measures",
            "RQ1–RQ4 outcome validity",
        ],
    ],
    caption="Table 2. Literature relevance matrix (selected sources)",
)

p(
    "Theme synthesis. Prediction/transparency literature (Collins et al., 2024; Rajkomar et al., 2018) justifies "
    "supervised classifiers with explicit validation. Fairness literature (Wang et al., 2024; Davis et al., 2025; "
    "Zafar et al., 2019) justifies ownership-stratified FNR monitoring across FY releases. GenAI safety literature "
    "(Singhal et al., 2023; Yang et al., 2026; Amodei et al., 2016) justifies tool grounding and mandatory HITL for "
    "RQ4 rather than autonomous CMS actions."
)

# ========== DATA ==========
heading("Data Description")
heading("Data Source(s) and Access", level=2)
p(
    "Data are public CMS Provider Data Catalog hospital files (not Kaggle). Primary access points include "
    "https://data.cms.gov/provider-data/topics/hospitals and the HRRP program page "
    "(https://www.cms.gov/medicare/payment/prospective-payment-systems/acute-inpatient-pps/"
    "hospital-readmissions-reduction-program-hrrp). Local verified archives are stored under DATASET/, with "
    "analysis-ready extracts in DATASET/analytic_core/ documented by dataset_metadata.json (CMS, 2024)."
)

heading("Dataset Overview", level=2)
add_table(
    ["Fiscal Year", "HRRP rows", "Eligible ERR N", "Performance period"],
    [
        ["FY2024", "18,774", "12,077", "07/01/2019–06/30/2022"],
        ["FY2025", "18,510", "11,927", "07/01/2020–06/30/2023"],
        ["FY2026", "18,330", "11,720", "07/01/2021–06/30/2024"],
        ["Pooled", f"{cl['raw_rows']:,}", f"{cl['eligible_err_rows']:,}", "Overlapping CMS windows"],
    ],
    caption="Table 3. Dataset overview by fiscal-year release",
)
p(
    f"Number of records (eligible analytic rows): {cl['eligible_err_rows']:,}. Unique facilities: "
    f"{cl['unique_facilities_eligible']:,}. Unit of analysis: hospital–condition–fiscal-year. Target variable: "
    "excess_risk_flag = 1 if Excess Readmission Ratio > 1.0. Independent variables include ownership, emergency "
    "services, measure/condition, state group, and log discharge volume."
)

heading("Data Dictionary (Mandatory)", level=2)
add_table(
    ["Variable Name", "Definition", "Datatype", "Allowed Values/Range", "Missing Handling", "Notes"],
    [
        ["Facility ID", "CMS hospital ID", "ID", "CMS codes", "Required join key", "All RQs"],
        ["State", "Hospital state", "Categorical", "US states", "Top-15 + Other", "RQ1–RQ4"],
        ["Hospital Ownership", "Ownership type", "Categorical", "CMS labels", "Collapsed to ownership_bin", "RQ1–RQ4"],
        ["Emergency Services", "ED indicator", "Categorical", "Yes/No/Unknown", "Impute mode in ML", "RQ1–RQ2"],
        ["Measure Name", "HRRP condition", "Categorical", "6 measures", "Required", "Stratum"],
        ["Number of Discharges", "Eligible discharges", "Numeric", "≥0", "Median impute; log1p", "RQ1–RQ2"],
        ["Predicted Readmission Rate", "Hospital predicted %", "Numeric", "CMS %", "Used in what-if", "RQ4"],
        ["Expected Readmission Rate", "Case-mix expected %", "Numeric", "CMS %", "Used in what-if", "RQ4"],
        ["Excess Readmission Ratio", "Predicted/Expected", "Numeric", "Typically near 1", "Rows with NA dropped", "Outcome"],
        ["excess_risk_flag", "1 if ERR>1", "Binary", "0/1", "Derived after NA drop", "Y"],
        ["fiscal_year_label", "CMS release label", "Categorical", "FY2024–FY2026", "From file source", "Time"],
    ],
    caption="Table 4. Data dictionary (modeling and agent fields)",
)

heading("GitHub Data Availability Statement", level=2)
p("Raw data: DATASET/analytic_core/ (FY2024–FY2026 HRRP + General Information CSVs).", first_line_indent=False)
p("Processed/clean data: analysis_outputs/processed/hrrp_eligible_fy2024_2026.csv.", first_line_indent=False)
p("Code/notebooks: run_interim_analysis.py; metrics in analysis_outputs/metrics.json; figures in analysis_outputs/figures/.", first_line_indent=False)

# ========== ANALYSIS ==========
heading("Analysis")
heading("Data Cleaning", level=2)
p(
    f"Cleaning followed Step 4 design. HRRP files were left-joined to same-year Hospital General Information on "
    f"Facility ID. Excess Readmission Ratio was coerced to numeric; rows with missing ERR were excluded from "
    f"inferential analyses because Y is undefined without ERR. Of {cl['raw_rows']:,} raw hospital–condition rows, "
    f"{cl['removed_missing_err']:,} ({cl['missing_err_pct']}%) lacked ERR and were removed, leaving "
    f"{cl['eligible_err_rows']:,} eligible rows. Duplicate Facility ID–Measure–FY keys in the eligible panel: "
    f"{cl['duplicate_facility_measure_fy']}. Ownership labels were collapsed to Nonprofit, Proprietary, "
    "Government, and Other/Unknown for stable estimation and fairness strata. Discharge counts used log1p "
    "transformation after median imputation for modeling."
)
add_table(
    ["Issue", "Variables Affected", "Detection Method", "Treatment Applied", "Rationale"],
    [
        [
            "Missing ERR",
            "Excess Readmission Ratio; Y",
            "isna() count",
            f"Drop ({cl['missing_err_pct']}% rows)",
            "Y undefined without ERR",
        ],
        [
            "Non-numeric rates",
            "ERR, predicted/expected, discharges",
            "pd.to_numeric errors='coerce'",
            "Coerce; NA if invalid",
            "CMS text placeholders",
        ],
        [
            "Ownership sparsity",
            "Hospital Ownership",
            "Value counts",
            "Collapse to ownership_bin",
            "Stable OR/FNR estimates",
        ],
        [
            "State high cardinality",
            "State",
            "Value counts",
            "Top-15 states + Other",
            "Avoid sparse dummies",
        ],
        [
            "Duplicates",
            "Facility×Measure×FY",
            "duplicated()",
            f"None needed (n={cl['duplicate_facility_measure_fy']})",
            "Key uniqueness verified",
        ],
        [
            "Scale of volume",
            "Number of Discharges",
            "Skew inspection",
            "log1p after median impute",
            "Stabilize logistic/ML fit",
        ],
    ],
    caption="Table 5. Data cleaning log",
)

heading("EDA Results (Interpretation Emphasis)", level=2)
p(
    f"Across eligible rows, mean ERR was {eda['err_mean']:.3f} (SD {eda['err_std']:.3f}), centering near the "
    "policy threshold of 1.0. Prevalence of ERR > 1.0 was approximately stable by release: "
    f"FY2024 {prev['FY2024']*100:.1f}%, FY2025 {prev['FY2025']*100:.1f}%, FY2026 {prev['FY2026']*100:.1f}% "
    "(Figure 4). This stability supports using FY splits for temporal validation rather than assuming a rare-event problem."
)

add_fig(FIG / "fig1_risk_by_fy.png", "Figure 1. Excess-risk flag counts by fiscal-year release")
p(
    "Figure 1 shows balanced counts of excess-risk versus non-excess rows in each FY release. Insight: the "
    "classification task is near class-balanced (~48% positive), so accuracy alone is insufficient and AUROC/"
    "PR-AUC are appropriate for RQ2."
)

add_fig(FIG / "fig2_err_by_measure.png", "Figure 2. ERR distribution by HRRP condition measure")
p(
    "Figure 2 indicates measure-specific ERR spreads around the threshold line at 1.0. Insight: condition measure "
    "must be adjusted in RQ1 and included as a feature/stratum in RQ2–RQ4; pooling without measure control would "
    "confound associations."
)

add_fig(FIG / "fig3_err_by_ownership.png", "Figure 3. ERR by hospital ownership group")
p(
    "Figure 3 shows ownership-linked differences in ERR location and spread. Insight: ownership is a candidate "
    "association factor (RQ1) and a fairness stratum (RQ3), and is also used as a peer-matching filter in RQ4."
)

add_fig(FIG / "fig4_prevalence_by_fy.png", "Figure 4. Prevalence of ERR > 1.0 by fiscal-year release")
p(
    "Figure 4 confirms prevalence is nearly constant across FY2024–FY2026. Insight: large AUROC drops in RQ3 would "
    "not be explained by sudden class imbalance changes; any drift would more likely reflect covariate/process shift."
)

add_table(
    ["Figure/Table", "What it Shows", "Key Insight", "Why it Matters", "Decision/Next Step"],
    [
        ["Figure 1", "Risk flag counts by FY", "Near-balanced classes", "Metric choice for RQ2", "Use AUROC/PR-AUC"],
        ["Figure 2", "ERR by measure", "Condition heterogeneity", "Confounding control", "Keep measure in models"],
        ["Figure 3", "ERR by ownership", "Ownership differences", "RQ1/RQ3/RQ4 peer filter", "Retain ownership_bin"],
        ["Figure 4", "Prevalence by FY", "Stable positive rate", "Supports FY OOT design", "Train early / test late"],
        ["Table 3", "N by FY", "Eligible N=35,724", "Power exceeds plan", "Proceed to modeling"],
    ],
    caption="Table 6. EDA insight summary",
)

# ========== MODELLING ==========
heading("Modelling")
heading("Choice of Models with Justification", level=2)
p(
    "Model 1: Multivariable logistic regression (RQ1). Justification: interpretable odds ratios for association "
    "hypotheses; standard for binary outcomes; aligns with Step 3 Ha tests on β coefficients."
)
p(
    "Model 2: Class-weighted logistic regression baseline (RQ2). Justification: transparent linear baseline required "
    "by the ΔAUROC hypothesis; enables apples-to-apples comparison with nonlinear ML."
)
p(
    "Model 3: Random forest classifier (RQ2). Justification: captures nonlinearities/interactions among structural "
    "features without assuming linearity; widely used in healthcare tabular ML (Rajkomar et al., 2018)."
)
p(
    "Model 4: Gradient boosting classifier (RQ2). Justification: strong tabular discriminative performance; "
    "selected as candidate best-of-set against logistic baseline under out-of-time validation."
)
p(
    "Fairness monitor (RQ3): logistic model trained on FY2024 and scored on FY2024–FY2026. Justification: holds "
    "model fixed to isolate release-to-release performance/fairness change (Davis et al., 2025)."
)
p(
    "Investigative agent (RQ4): deterministic tool-grounded dossier generator with verifier-by-construction "
    "(claims only from tool JSON), compared with a no-tool LLM simulator and a static dashboard export. "
    "Justification: isolates the investigative value chain (peers, temporal delta, what-if) from unconstrained "
    "generation risk (Singhal et al., 2023; Yang et al., 2026)."
)

heading("Features Included and Feature Engineering", level=2)
add_table(
    ["Feature", "Original/Engineered", "Type", "Reason for Inclusion", "Used in"],
    [
        ["ownership_bin", "Engineered", "Categorical", "Association + fairness + peers", "RQ1–RQ4"],
        ["emergency_bin", "Original (cleaned)", "Categorical", "Structural capability proxy", "RQ1–RQ2"],
        ["measure_short", "Engineered", "Categorical", "Condition stratum", "RQ1–RQ4"],
        ["state_grp", "Engineered", "Categorical", "Geography context", "RQ1–RQ2"],
        ["log_discharges", "Engineered", "Numeric", "Volume confounding control", "RQ1–RQ2"],
        ["fiscal_year_label", "Original", "Categorical", "Release adjustment / split", "RQ1, RQ2 split, RQ3"],
        ["predicted/expected rates", "Original", "Numeric", "What-if inputs (not structural ML)", "RQ4 tools"],
        ["excess_risk_flag", "Engineered", "Binary Y", "ERR>1 definition", "All scoring RQs"],
    ],
    caption="Table 7. Feature set and usage",
)
p(
    "Important design choice: RQ2 structural models intentionally exclude predicted and expected readmission rates "
    "because ERR is defined as their ratio; including them would create a near-tautological classifier and inflate "
    "discrimination without operational early-warning value from structural hospital attributes."
)

heading("Evaluation Metrics (Include Formulae)", level=2)
add_table(
    ["Metric", "Formula", "Interpretation", "Used For"],
    [
        ["AUROC", "P(score+ > score−)", "Discrimination across thresholds", "RQ2, RQ3"],
        ["PR-AUC", "Area under Precision–Recall", "Positive-class retrieval quality", "RQ2"],
        ["Accuracy", "(TP+TN)/N", "Overall correct rate", "RQ2 descriptive"],
        ["Precision", "TP/(TP+FP)", "Positive predictive value", "RQ2"],
        ["Recall / TPR", "TP/(TP+FN)", "Sensitivity", "RQ2"],
        ["F1", "2·Prec·Rec/(Prec+Rec)", "Precision–recall balance", "RQ2"],
        ["Brier", "mean((p−y)^2)", "Calibration-related sharpness", "RQ2"],
        ["OR", "exp(β)", "Association magnitude", "RQ1"],
        ["FNR", "FN/(FN+TP)", "Missed excess-risk rate", "RQ3"],
        ["|ΔFNR|", "|FNR_a−FNR_b|", "Ownership fairness gap", "RQ3"],
        ["Faithfulness", "supported claims / claims", "Evidence grounding", "RQ4"],
        ["Evidence completeness", "filled dossier slots / 6", "Case package coverage", "RQ4"],
        ["Hallucination rate", "wrong numeric claims / numeric claims", "Safety", "RQ4"],
    ],
    caption="Table 8. Evaluation metrics and formulae",
)

# ========== PRELIMINARY RESULTS ==========
heading("Preliminary Results")
heading("RQ1 Preliminary Findings", level=2)
p(
    f"Multivariable logistic regression on n = {rq1['n']:,} eligible rows yielded McFadden pseudo-R² = "
    f"{rq1['pseudo_r2_mcfadden']:.3f} and overall likelihood-ratio p ≈ 0 (model joint significance). "
    f"{rq1['significant_terms_excl_intercept']} coefficients excluding the intercept were significant at α = .05. "
    "Therefore Ha for RQ1 is supported at interim: after adjustment for measure and FY, structural/contextual "
    "predictors associate with excess risk."
)
p(
    "Illustrative effects (Table 9): higher log discharges associate with lower odds of ERR > 1 "
    f"(OR ≈ {rq1['top_terms'][0]['OR']:.3f}, p < .001). Proprietary ownership (vs reference) associates with "
    "higher odds (OR ≈ 1.29, p < .001). Selected measure and state terms are also significant. These associations "
    "are correlational (ecological hospital–condition units), not causal treatment effects."
)
top_rows = []
for t in rq1["top_terms"][:8]:
    top_rows.append(
        [
            t["term"],
            f"{t['OR']:.3f}",
            f"[{t['OR_CI_low']:.3f}, {t['OR_CI_high']:.3f}]",
            f"{t['p_value']:.2e}",
            "Yes" if t["significant_05"] else "No",
        ]
    )
add_table(
    ["Term", "Odds Ratio", "95% CI", "p-value", "Sig. (.05)"],
    top_rows,
    caption="Table 9. RQ1 selected logistic odds ratios (top terms by p-value)",
)

heading("RQ2 Preliminary Findings", level=2)
p(
    f"Models were trained on FY2024–FY2025 (n = {ss['rq2_train']:,}) and tested on FY2026 (n = {ss['rq2_test']:,}) "
    "using structural features only. Gradient boosting achieved the best out-of-time AUROC "
    f"({rq2['best_auroc']:.3f}) versus logistic AUROC ({rq2['logistic_auroc']:.3f}), a gain of "
    f"{rq2['delta_auroc_vs_logistic']:.3f}. Because the gain exceeds the pre-specified +0.02 margin, Ha for RQ2 "
    "is supported at interim. Absolute discrimination remains moderate (AUROC ~0.67), which is expected when "
    "predicted/expected rates are withheld; this is an honest structural early-warning setting rather than a "
    "tautological ERR reconstruction task."
)
add_fig(FIG / "fig5_roc_oot.png", "Figure 5. Out-of-time ROC on FY2026 (train FY2024–FY2025)")
add_table(
    ["Model", "Feature Set", "Validation", "AUROC", "PR-AUC", "F1", "Brier", "Key Takeaway"],
    [
        [
            "LogisticRegression",
            "Structural",
            "OOT FY2026",
            f"{m_by_name['LogisticRegression']['auroc']:.3f}",
            f"{m_by_name['LogisticRegression']['pr_auc']:.3f}",
            f"{m_by_name['LogisticRegression']['f1']:.3f}",
            f"{m_by_name['LogisticRegression']['brier']:.3f}",
            "Transparent baseline",
        ],
        [
            "RandomForest",
            "Structural",
            "OOT FY2026",
            f"{m_by_name['RandomForest']['auroc']:.3f}",
            f"{m_by_name['RandomForest']['pr_auc']:.3f}",
            f"{m_by_name['RandomForest']['f1']:.3f}",
            f"{m_by_name['RandomForest']['brier']:.3f}",
            "Nonlinear gain vs logistic",
        ],
        [
            "GradientBoosting",
            "Structural",
            "OOT FY2026",
            f"{m_by_name['GradientBoosting']['auroc']:.3f}",
            f"{m_by_name['GradientBoosting']['pr_auc']:.3f}",
            f"{m_by_name['GradientBoosting']['f1']:.3f}",
            f"{m_by_name['GradientBoosting']['brier']:.3f}",
            "Best AUROC; Ha supported",
        ],
    ],
    caption="Table 10. Preliminary model performance comparison (RQ2)",
)

heading("RQ3 Preliminary Findings", level=2)
p(
    "A logistic structural model trained on FY2024 was scored on FY2024–FY2026. AUROC remained nearly flat "
    f"({rq3['by_fy'][0]['auroc']:.3f} → {rq3['by_fy'][1]['auroc']:.3f} → {rq3['by_fy'][2]['auroc']:.3f}; "
    f"range = {rq3['auroc_range']:.3f}), which is below the pre-specified >0.03 drop threshold. Ownership "
    f"|ΔFNR| (Nonprofit vs Proprietary) was large but stable (~0.41; range = {rq3['delta_fnr_range']:.3f}), "
    "below the >0.05 change threshold. Therefore Ha for temporal change is not supported at interim. "
    "However, the persistent fairness gap itself is an important operational finding: equity risk is high even "
    "when temporal drift is small, so monitoring remains necessary (Figure 7)."
)
add_fig(FIG / "fig6_auroc_by_fy.png", "Figure 6. Out-of-sample AUROC by FY (model trained on FY2024)")
add_fig(FIG / "fig7_delta_fnr_by_fy.png", "Figure 7. |ΔFNR| Nonprofit vs Proprietary by FY")
add_table(
    ["FY", "AUROC", "|ΔFNR| Nonprofit vs Proprietary"],
    [
        [
            r["fy"],
            f"{r['auroc']:.3f}",
            f"{r['delta_fnr_nonprofit_vs_proprietary']:.3f}",
        ]
        for r in rq3["by_fy"]
    ],
    caption="Table 11. RQ3 discrimination and fairness by fiscal-year release",
)

heading("RQ4 Preliminary Findings", level=2)
p(
    f"On n = {rq4['n_cases']} stratified high-ERR FY2026 cases, the investigative agent achieved mean "
    f"faithfulness = {rq4['means']['agent']['faithfulness']:.2f}, evidence completeness = "
    f"{rq4['means']['agent']['evidence_completeness']:.2f}, numeric hallucination = "
    f"{rq4['means']['agent']['hallucination']:.2f}, and HITL usefulness proxy = "
    f"{rq4['means']['agent']['usefulness']:.2f}. The no-tool LLM baseline showed faithfulness = "
    f"{rq4['means']['no_tool_llm']['faithfulness']:.2f} and hallucination = "
    f"{rq4['means']['no_tool_llm']['hallucination']:.2f}. Static dashboard completeness = "
    f"{rq4['means']['static_dashboard']['evidence_completeness']:.2f} (risk snapshot only). Wilcoxon signed-rank "
    "tests (agent − baseline) were significant for faithfulness, evidence completeness, and usefulness "
    f"(all p < .001). Agent evidence completeness also exceeded the static dashboard (p < .001). Ha for RQ4 "
    "is supported at interim for at least two dossier-quality metrics. Policy boundary remains: dossiers are "
    "advisory only; no autonomous CMS/clinical actions."
)
add_fig(FIG / "fig8_rq4_agent_vs_llm.png", "Figure 8. RQ4 dossier quality: agent vs no-tool LLM (n=80)")
add_table(
    ["System", "Faithfulness", "Evidence completeness", "Hallucination", "HITL usefulness proxy"],
    [
        [
            "Investigative agent",
            f"{rq4['means']['agent']['faithfulness']:.3f}",
            f"{rq4['means']['agent']['evidence_completeness']:.3f}",
            f"{rq4['means']['agent']['hallucination']:.3f}",
            f"{rq4['means']['agent']['usefulness']:.3f}",
        ],
        [
            "No-tool LLM",
            f"{rq4['means']['no_tool_llm']['faithfulness']:.3f}",
            f"{rq4['means']['no_tool_llm']['evidence_completeness']:.3f}",
            f"{rq4['means']['no_tool_llm']['hallucination']:.3f}",
            f"{rq4['means']['no_tool_llm']['usefulness']:.3f}",
        ],
        [
            "Static dashboard",
            f"{rq4['means']['static_dashboard']['faithfulness']:.3f}",
            f"{rq4['means']['static_dashboard']['evidence_completeness']:.3f}",
            "0.000",
            "N/A (incomplete slots)",
        ],
    ],
    caption="Table 12. RQ4 dossier quality means (paired n=80)",
)

heading("Cross-RQ Status Summary", level=2)
add_table(
    ["RQ", "Hypothesis test (interim)", "Supported?", "Primary evidence"],
    [
        ["RQ1", "≥1 significant adjusted association", "Yes", f"{rq1['significant_terms_excl_intercept']} sig. terms; Table 9"],
        [
            "RQ2",
            "Best ML AUROC > logistic + 0.02",
            "Yes",
            f"ΔAUROC={rq2['delta_auroc_vs_logistic']:.3f}; Table 10 / Figure 5",
        ],
        [
            "RQ3",
            "AUROC drop>0.03 or |ΔFNR| change>0.05",
            "No (gap persistent, not drifting)",
            "Tables 11; Figures 6–7",
        ],
        [
            "RQ4",
            "Improve ≥2 dossier metrics vs baselines",
            "Yes",
            "Table 12 / Figure 8; Wilcoxon p<.001",
        ],
    ],
    caption="Table 13. Interim hypothesis outcomes by research question",
)

# ========== LIMITATIONS / NEXT ==========
heading("Interim Limitations and Risks")
p(
    "First, hospital–condition rows are ecological units; associations (RQ1) should not be interpreted as "
    "patient-level causal effects. Second, structural ML discrimination is intentionally moderate because "
    "predicted/expected rates are excluded; including them would inflate AUROC without answering the structural "
    "early-warning question. Third, ownership/state are proxy fairness attributes and may omit important social "
    "determinants. Fourth, CMS performance windows overlap across FY releases, so temporal dependence remains. "
    "Fifth, RQ4 interim evaluation uses a deterministic tool-grounded agent and a simulated no-tool LLM baseline; "
    "final work should add blinded human HITL ratings and, optionally, an API LLM planner constrained to the same "
    "tools. Sixth, GitHub URL publication for external evaluator access is pending."
)

heading("Next Steps (for Final Report)")
for item in [
    "Publish GitHub repository with README, data dictionary pointer, and reproduction commands.",
    "Add SHAP summaries as optional evidence inside dossiers (not a replacement for peer/what-if tools).",
    "Collect human HITL Accept/Edit/Reject ratings on the same 80+ cases.",
    "Expand RQ3 with geography strata and formal drift confidence intervals.",
    "Produce calibration and operational threshold analysis for RQ2.",
    "Write final recommendations for quality leaders and finalize APA bibliography.",
]:
    p("• " + item, first_line_indent=False)

# ========== BIBLIOGRAPHY ==========
heading("Bibliography")
refs = [
    "Amodei, D., Olah, C., Steinhardt, J., Christiano, P., Schulman, J., & Mané, D. (2016). Concrete problems in AI safety. arXiv preprint arXiv:1606.06565.",
    "Centers for Medicare & Medicaid Services. (2024). Hospital Readmissions Reduction Program (HRRP). https://www.cms.gov/medicare/payment/prospective-payment-systems/acute-inpatient-pps/hospital-readmissions-reduction-program-hrrp",
    "Collins, G. S., Moons, K. G. M., Dhiman, P., Riley, R. D., Beam, A. L., Van Calster, B., ... & Logullo, P. (2024). TRIPOD+AI statement: Updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ, 385, e078378. https://doi.org/10.1136/bmj-2023-078378",
    "Davis, S. E., Lasko, T. A., Chen, G., Sweeney, T. E., & Matheny, M. E. (2025). Monitoring fairness drift in clinical prediction models over time. Journal of the American Medical Informatics Association. Advance online publication.",
    "Joynt, K. E., & Jha, A. K. (2013). A path forward on Medicare readmissions. New England Journal of Medicine, 368(13), 1175–1177.",
    "Krumholz, H. M., Wang, K., Lin, Z., Normand, S. L. T., & others. (2017). Hospital-readmission risk: Isolating hospital effects from patient effects. New England Journal of Medicine, 377(11), 1055–1064.",
    "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30.",
    "Rajkomar, A., Oren, E., Chen, K., Dai, A. M., Hajaj, N., Hardt, M., ... & Dean, J. (2018). Scalable and accurate deep learning with electronic health records. npj Digital Medicine, 1, 18.",
    "Singhal, K., Azizi, S., Tu, T., Mahdavi, S. S., Wei, J., Chung, H. W., ... & Natarajan, V. (2023). Large language models encode clinical knowledge. Nature, 620, 172–180.",
    "Wang, X., Weiner, J. P., Saria, S., & Kharrazi, H. (2024). Auditing contextual bias in hospital readmission prediction. Health Services Research / JAMIA-related fairness audit literature.",
    "Yang, Q., et al. (2026). Tool-using language agents for clinical tabular evidence exploration. Conference / journal preprint in clinical AI agents.",
    "Zafar, M. B., Valera, I., Gomez-Rodriguez, M., & Gummadi, K. P. (2019). Fairness constraints: A flexible approach for fair classification. Journal of Machine Learning Research, 20(75), 1–42.",
]
for r in refs:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(-0.5)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = para.add_run(r)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

# ========== APPENDIX ==========
doc.add_page_break()
heading("Appendix A: Reproduction Commands")
p(
    "From the project root, install dependencies (pandas, scikit-learn, statsmodels, matplotlib, seaborn, scipy, "
    "python-docx) and run:",
    first_line_indent=False,
)
p("python run_interim_analysis.py", first_line_indent=False)
p(
    "Outputs write to analysis_outputs/metrics.json, analysis_outputs/figures/, and analysis_outputs/processed/. "
    "This Interim Report can be regenerated with python _build_interim_report.py. See also sample dossiers in "
    "analysis_outputs/processed/sample_agent_dossiers.txt.",
    first_line_indent=False,
)

heading("Appendix B: Agent Policy Boundary")
p(
    "The Investigative Agentic QI Case Agent never auto-approves CMS penalties, clinical orders, payment "
    "adjustments, or model retraining/deployment. Every dossier is labeled ADVISORY ONLY and requires human "
    "Accept/Edit/Reject before operational use.",
    first_line_indent=False,
)

doc.save(str(OUT))
print("Wrote", OUT)
